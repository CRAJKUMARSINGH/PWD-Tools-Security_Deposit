"""
Convert Excel security refund sheet to Word document format
"""
import openpyxl
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_border(cell, **kwargs):
    """
    Set cell borders in Word table
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create borders element
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), 'single')
            edge_element.set(qn('w:sz'), '4')
            edge_element.set(qn('w:space'), '0')
            edge_element.set(qn('w:color'), '000000')
            tcBorders.append(edge_element)
    
    tcPr.append(tcBorders)

def convert_excel_to_word(excel_file, word_file):
    """
    Convert Excel sheet to Word document preserving formatting
    """
    print(f"Loading Excel file: {excel_file}")
    wb = openpyxl.load_workbook(excel_file)
    ws = wb.active
    
    # Create Word document
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Get dimensions
    max_row = ws.max_row
    max_col = ws.max_column
    
    print(f"Creating table with {max_row} rows and {max_col} columns")
    
    # Create table in Word
    table = doc.add_table(rows=max_row, cols=max_col)
    table.style = 'Table Grid'
    
    # Copy data and formatting from Excel to Word
    for row_idx, row in enumerate(ws.iter_rows(min_row=1, max_row=max_row), start=0):
        for col_idx, cell in enumerate(row):
            word_cell = table.rows[row_idx].cells[col_idx]
            
            # Copy cell value
            if cell.value:
                word_cell.text = str(cell.value)
            
            # Apply formatting
            if word_cell.text:
                paragraph = word_cell.paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(word_cell.text)
                
                # Font formatting
                if cell.font:
                    if cell.font.bold:
                        run.font.bold = True
                    if cell.font.size:
                        run.font.size = Pt(cell.font.size / 1.33)  # Convert Excel points
                    if cell.font.color and cell.font.color.rgb:
                        rgb = cell.font.color.rgb
                        run.font.color.rgb = RGBColor(int(rgb[2:4], 16), int(rgb[4:6], 16), int(rgb[6:8], 16))
                
                # Alignment
                if cell.alignment:
                    if cell.alignment.horizontal == 'center':
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif cell.alignment.horizontal == 'right':
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif cell.alignment.horizontal == 'left':
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add borders
            set_cell_border(word_cell, top=True, left=True, bottom=True, right=True)
    
    # Save Word document
    print(f"Saving Word document: {word_file}")
    doc.save(word_file)
    print("Conversion completed successfully!")

if __name__ == "__main__":
    excel_file = "Blank_Security_Refund_Template.xlsx"
    word_file = "Blank_Security_Refund_Template.docx"
    
    if os.path.exists(excel_file):
        convert_excel_to_word(excel_file, word_file)
        print(f"\n✓ Word document created!")
        print(f"✓ Output file: {word_file}")
    else:
        print(f"Error: Excel file not found: {excel_file}")
